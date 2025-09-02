from abc import ABC, abstractmethod
from docx import Document
from PyPDF2 import PdfReader


class FileContentGetter(ABC):
    @abstractmethod
    def get_content(self, file_path: str) -> str:
        pass


class FileContentGetterFactory:
    """
    Factory class for creating FileContentGetter objects
    """
    @staticmethod
    def create(file_path: str) -> FileContentGetter:
        """
        Create a FileContentGetter object based on the file extension
        :param file_path: file path
        """
        if file_path.endswith('.docx'):
            return DocxContentGetter()
        elif file_path.endswith('.pdf'):
            return PdfContentGetter()
        elif file_path.endswith('.txt'):
            return TxtContentGetter()
        else:
            raise ValueError("Unsupported file type")


class DocxContentGetter(FileContentGetter):
    """
    Get content from a docx file
    """
    def get_content(self, file_path: str) -> str:
        """
        加载docx文件文本内容
        :param file_path: 文件路径
        """
        # 打开并读取docx文件，尝试转换为md标记字符串，移交md_parser进一步解析分割。
        document = Document(file_path)
        markdown_text = ""
        if document.paragraphs:
            title = document.paragraphs[0].text.strip()
            markdown_text += f"# {title}\n\n"
        for para in document.paragraphs[1:]:  # 跳过第一个段落
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])
                markdown_text += f"{'#' * level} {para.text}\n\n"
            else:
                markdown_text += f"{para.text}\n\n"
        return markdown_text


class PdfContentGetter(FileContentGetter):
    """
    Get content from a pdf file
    """
    def get_content(self, file_path: str) -> str:
        """
        加载pdf文件文本内容
        :param file_path: 文件路径
        """
        # 读取pdf文件内容
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            # 构建字符串列表，存储PDF文件中的文本内容
            text_list = []
            # 获取PDF文件中的总页数
            num_pages = len(reader.pages)

            # 逐页读取文本内容
            for page_num in range(num_pages):
                # 获取当前页的文本内容
                page = reader.pages[page_num]
                # 将当前页的文本内容添加到text_list列表中
                text_list.append(page.extract_text())
            # 将text_list列表中的文本内容合并为一个字符串
            text = "\n".join(text_list)
            return text

    def _get_content(self, file_path):
        """
        模型pdf2md，对设备性能要求极高
        """
        # from marker.convert import convert_single_pdf
        # from marker.models import load_all_models
        # model_list = load_all_models()
        #
        # input_pdf = r"C:\Users\18468\Desktop\unigraph_test.pdf"
        # full_text, images, out_meta = convert_single_pdf(input_pdf, model_list, 3)
        #
        # parsed_headings = DocxStructuredContentGetter.parse_markdown_headings(markdown_text=full_text)
        # rdf_triples = DocxStructuredContentGetter.generate_kg_triples(parsed_headings=parsed_headings)
        # structured_kg = DocxStructuredContentGetter.convert_to_custom_triples(triples=rdf_triples)
        # md_dic = DocxStructuredContentGetter.parse_markdown2dic(markdown_text=full_text)
        # return md_dic, structured_kg
        pass


class TxtContentGetter(FileContentGetter):
    """
    Get content from a txt file
    """
    def get_content(self, file_path: str) -> str:
        """
        加载txt文件文本内容
        :param file_path: 文件路径
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
